"""
DDR Report Generator - AI-Powered Detailed Diagnostic Report System
Uses Google Gemini (free tier) to process inspection + thermal reports
and generate structured, client-ready DDR reports with images.
"""

import os
import io
import json
import base64
import tempfile
import re
import uuid
from datetime import datetime
from pathlib import Path

from flask import Flask, request, render_template, send_file, jsonify, url_for
from werkzeug.utils import secure_filename

import fitz  # PyMuPDF
from PIL import Image
try:
    from google import genai as genai_new
    USE_NEW_SDK = True
except ImportError:
    import google.generativeai as genai
    USE_NEW_SDK = False
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak, KeepTogether
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB
app.config['UPLOAD_FOLDER'] = tempfile.mkdtemp()
app.config['OUTPUT_FOLDER'] = tempfile.mkdtemp()

ALLOWED_EXTENSIONS = {'pdf'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ─────────────────────────────────────────────
# PDF PROCESSING - Extract text + images
# ─────────────────────────────────────────────

def extract_from_pdf(pdf_path):
    """Extract text and images from a PDF file using PyMuPDF."""
    doc = fitz.open(pdf_path)
    pages_data = []
    all_images = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        
        page_images = []
        image_list = page.get_images(full=True)
        
        page_img_count = 0
        for img_idx, img_info in enumerate(image_list):
            if page_img_count >= 4:  # Max 4 significant images per page
                break
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                if base_image and base_image.get("image"):
                    img_bytes = base_image["image"]
                    img_ext = base_image.get("ext", "png")
                    
                    # Filter out small images (logos, icons, UI elements)
                    try:
                        pil_img = Image.open(io.BytesIO(img_bytes))
                        w, h = pil_img.size
                        # Skip tiny images, icons, and very small thumbnails
                        if w < 150 or h < 150:
                            continue
                        # Skip images smaller than 5KB (likely UI elements)
                        if len(img_bytes) < 5000:
                            continue
                    except:
                        continue
                    
                    img_id = f"page{page_num+1}_img{img_idx+1}"
                    img_b64 = base64.b64encode(img_bytes).decode('utf-8')
                    
                    page_images.append({
                        'id': img_id,
                        'data': img_bytes,
                        'b64': img_b64,
                        'ext': img_ext,
                        'width': w,
                        'height': h,
                        'page': page_num + 1
                    })
                    page_img_count += 1
            except Exception as e:
                continue
        
        pages_data.append({
            'page_num': page_num + 1,
            'text': text,
            'images': page_images
        })
        all_images.extend(page_images)
    
    doc.close()
    
    full_text = "\n\n".join([f"--- Page {p['page_num']} ---\n{p['text']}" for p in pages_data])
    return full_text, all_images, pages_data


def classify_images(inspection_images, thermal_images):
    """Classify images into categories based on source document."""
    classified = {
        'inspection_photos': [],
        'thermal_images': []
    }
    
    for img in inspection_images:
        classified['inspection_photos'].append(img)
    
    for img in thermal_images:
        # Thermal images from the thermal report
        classified['thermal_images'].append(img)
    
    return classified


# ─────────────────────────────────────────────
# AI PROCESSING - Gemini API
# ─────────────────────────────────────────────

DDR_PROMPT = """You are an expert property inspection analyst. You must analyze the provided inspection report and thermal imaging report to generate a comprehensive Detailed Diagnostic Report (DDR).

INSPECTION REPORT TEXT:
{inspection_text}

THERMAL REPORT TEXT:
{thermal_text}

INSPECTION REPORT has {num_inspection_images} photographs.
THERMAL REPORT has {num_thermal_images} thermal images (each page has a thermal image + corresponding visual photo with temperature readings).

Based on the above data, generate a DDR report in the following JSON structure. Be thorough, accurate, and use ONLY information present in the documents. Do NOT invent facts.

Return ONLY valid JSON (no markdown, no backticks):

{{
    "report_metadata": {{
        "report_id": "auto-generated ID",
        "inspection_date": "date from report",
        "inspected_by": "names from report",
        "prepared_for": "client details from report",
        "property_type": "type from report",
        "property_address": "address if available or Not Available",
        "property_age": "age if available or Not Available",
        "floors": "number of floors",
        "previous_structural_audit": "Yes/No",
        "previous_repair_work": "Yes/No"
    }},
    "property_issue_summary": "A comprehensive 3-5 paragraph summary of ALL issues found across both reports. Mention dampness, leakage, tile hollowness, plumbing issues, external wall cracks, thermal anomalies etc. Reference specific flat numbers and areas.",
    
    "area_wise_observations": [
        {{
            "area_id": 1,
            "area_name": "e.g., Hall - Flat No. 103",
            "negative_side": {{
                "description": "What damage/issue is observed on the affected side",
                "details": "Detailed observation with specifics",
                "inspection_photo_refs": ["Photo numbers from inspection report e.g. Photo 1, Photo 2"],
                "thermal_image_refs": ["Thermal page numbers if applicable e.g. Thermal Page 1"]
            }},
            "positive_side": {{
                "description": "Source/cause side observation",
                "details": "Detailed observation about the source",
                "inspection_photo_refs": ["Photo numbers"],
                "thermal_image_refs": []
            }},
            "thermal_findings": {{
                "hotspot_temp": "temperature if available",
                "coldspot_temp": "temperature if available",
                "temperature_differential": "calculated difference",
                "interpretation": "What the thermal data indicates about moisture/leakage"
            }}
        }}
    ],
    
    "probable_root_causes": [
        {{
            "cause": "e.g., Gaps in bathroom tile joints",
            "explanation": "Detailed explanation of how this causes the observed damage",
            "affected_areas": ["List of areas affected by this cause"]
        }}
    ],
    
    "severity_assessment": [
        {{
            "area": "area name",
            "severity": "Critical/High/Moderate/Low",
            "reasoning": "Why this severity level was assigned",
            "urgency": "Immediate/Short-term/Long-term"
        }}
    ],
    
    "recommended_actions": [
        {{
            "priority": 1,
            "action": "Specific action to take",
            "description": "Detailed description of the repair/treatment",
            "applicable_areas": ["areas where this applies"],
            "estimated_urgency": "Immediate/Within 1 month/Within 3 months/Within 6 months"
        }}
    ],
    
    "additional_notes": [
        "Any additional observations, warnings, or notes"
    ],
    
    "missing_or_unclear_information": [
        "List any information that was missing, unclear, or marked N/A in the reports"
    ],
    
    "checklist_summary": {{
        "wc_checklist": {{
            "leakage_condition": "from report",
            "leakage_timing": "from report",
            "concealed_plumbing_issue": "Yes/No",
            "nahani_trap_damage": "Yes/No",
            "tile_joint_gaps": "Yes/No",
            "nahani_trap_gaps": "Yes/No",
            "tiles_broken": "Yes/No",
            "loose_plumbing_joints": "Yes/No"
        }},
        "external_wall_checklist": {{
            "interior_leakage": "from report",
            "leakage_timing": "from report",
            "concealed_plumbing_issue": "Yes/No",
            "internal_leakage_observed": "Yes/No",
            "cracks_on_external": "condition",
            "plumbing_pipes_condition": "condition",
            "algae_fungus_moss": "condition",
            "rcc_cracks": "condition"
        }}
    }},
    
    "thermal_analysis_summary": [
        {{
            "thermal_page": 1,
            "image_file": "filename",
            "hotspot": "temp",
            "coldspot": "temp",
            "differential": "temp diff",
            "area_mapped_to": "which inspection area this corresponds to",
            "finding": "interpretation"
        }}
    ],
    
    "summary_table": [
        {{
            "point_no": 1,
            "impacted_area_negative": "description of negative side",
            "point_no_positive": "1.1",
            "exposed_area_positive": "description of positive/source side"
        }}
    ]
}}

IMPORTANT RULES:
1. Extract ALL observations from BOTH documents
2. Do NOT invent information not present in the documents
3. If information is missing, write "Not Available"
4. If information conflicts between documents, mention the conflict
5. Map thermal images to corresponding inspection areas based on location/description
6. Use simple, client-friendly language
7. Include ALL 7 impacted areas from the inspection report
8. Include ALL thermal readings from the thermal report
9. Return ONLY valid JSON"""


def generate_ddr_with_gemini(api_key, inspection_text, thermal_text, num_inspection_images, num_thermal_images):
    """Use Google Gemini to generate the DDR report."""
    
    prompt = DDR_PROMPT.format(
        inspection_text=inspection_text,
        thermal_text=thermal_text,
        num_inspection_images=num_inspection_images,
        num_thermal_images=num_thermal_images
    )
    
    if USE_NEW_SDK:
        client = genai_new.Client(api_key=api_key)
        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt,
            config={
                'temperature': 0.1,
                'max_output_tokens': 8192,
            }
        )
        text = response.text.strip()
    else:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
        response = model.generate_content(
            prompt,
            generation_config=genai.GenerationConfig(
                temperature=0.1,
                max_output_tokens=8192,
            )
        )
        text = response.text.strip()
    
    # Clean markdown code fences if present
    if text.startswith('```'):
        text = re.sub(r'^```(?:json)?\n?', '', text)
        text = re.sub(r'\n?```$', '', text)
    
    return json.loads(text)


# ─────────────────────────────────────────────
# REPORT GENERATION - DOCX
# ─────────────────────────────────────────────

def save_images_to_temp(images, folder):
    """Save extracted images to temp folder and return paths."""
    paths = {}
    for img in images:
        img_path = os.path.join(folder, f"{img['id']}.{img['ext']}")
        with open(img_path, 'wb') as f:
            f.write(img['data'])
        paths[img['id']] = img_path
    return paths


def generate_docx(ddr_data, inspection_images, thermal_images, img_folder, output_path):
    """Generate a professional DOCX report from DDR data."""
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ---- COVER PAGE ----
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('DETAILED DIAGNOSTIC REPORT')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(139, 0, 0)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = ddr_data.get('report_metadata', {})
    run = subtitle.add_run(f"\n{meta.get('prepared_for', 'Property Inspection Report')}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(80, 80, 80)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"\nInspection Date: {meta.get('inspection_date', 'N/A')}")
    run.font.size = Pt(12)
    run = info.add_run(f"\nInspected By: {meta.get('inspected_by', 'N/A')}")
    run.font.size = Pt(12)
    run = info.add_run(f"\nReport ID: {meta.get('report_id', 'DDR-' + str(uuid.uuid4())[:8].upper())}")
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ---- DISCLAIMER ----
    doc.add_heading('Data and Information Disclaimer', level=1)
    disclaimer_text = (
        "This property inspection is not an exhaustive inspection of the structure, systems, or components. "
        "The inspection may not reveal all deficiencies. A health checkup helps to reduce some of the risk "
        "involved in the property/structure & premises, but it cannot eliminate these risks, nor can the "
        "inspection anticipate future events or changes in performance due to changes in use or occupancy. "
        "The inspection report may address issues that are code-based; however, this is NOT a code compliance "
        "inspection and does NOT verify compliance with manufacturer's installation instructions."
    )
    p = doc.add_paragraph(disclaimer_text)
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_page_break()
    
    # ---- TABLE OF CONTENTS ----
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        'Section 1: Property Information',
        'Section 2: Property Issue Summary',
        'Section 3: Area-wise Observations',
        'Section 4: Probable Root Causes',
        'Section 5: Severity Assessment',
        'Section 6: Recommended Actions',
        'Section 7: Thermal Analysis',
        'Section 8: Summary Table',
        'Section 9: Additional Notes',
        'Section 10: Missing/Unclear Information',
        'Appendix: Inspection Photographs',
        'Appendix: Thermal Images'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # ---- SECTION 1: PROPERTY INFORMATION ----
    doc.add_heading('Section 1: Property Information', level=1)
    
    info_table = doc.add_table(rows=9, cols=2)
    info_table.style = 'Light Grid Accent 1'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    fields = [
        ('Property Type', meta.get('property_type', 'N/A')),
        ('Property Address', meta.get('property_address', 'N/A')),
        ('Property Age', meta.get('property_age', 'N/A')),
        ('Floors', meta.get('floors', 'N/A')),
        ('Inspection Date', meta.get('inspection_date', 'N/A')),
        ('Inspected By', meta.get('inspected_by', 'N/A')),
        ('Previous Structural Audit', meta.get('previous_structural_audit', 'N/A')),
        ('Previous Repair Work', meta.get('previous_repair_work', 'N/A')),
        ('Report ID', meta.get('report_id', 'N/A')),
    ]
    
    for i, (label, value) in enumerate(fields):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = str(value)
        for cell in [info_table.cell(i, 0), info_table.cell(i, 1)]:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
    
    doc.add_paragraph()
    
    # ---- SECTION 2: ISSUE SUMMARY ----
    doc.add_heading('Section 2: Property Issue Summary', level=1)
    doc.add_paragraph(ddr_data.get('property_issue_summary', 'Not Available'))
    
    # ---- SECTION 3: AREA-WISE OBSERVATIONS ----
    doc.add_page_break()
    doc.add_heading('Section 3: Area-wise Observations', level=1)
    
    insp_img_paths = save_images_to_temp(inspection_images, img_folder)
    therm_img_paths = save_images_to_temp(thermal_images, img_folder)
    
    for obs in ddr_data.get('area_wise_observations', []):
        doc.add_heading(f"Area {obs.get('area_id', '')}: {obs.get('area_name', 'Unknown')}", level=2)
        
        # Negative side
        neg = obs.get('negative_side', {})
        doc.add_heading('Negative Side (Affected Area)', level=3)
        doc.add_paragraph(f"Description: {neg.get('description', 'N/A')}")
        doc.add_paragraph(f"Details: {neg.get('details', 'N/A')}")
        
        # Add relevant inspection photos
        photo_refs = neg.get('inspection_photo_refs', [])
        if photo_refs:
            doc.add_paragraph('Inspection Photographs:', style='Intense Quote')
            added = 0
            for img_id, img_path in insp_img_paths.items():
                if added >= 3:
                    break
                try:
                    doc.add_picture(img_path, width=Inches(2.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    added += 1
                except:
                    continue
        
        # Positive side
        pos = obs.get('positive_side', {})
        doc.add_heading('Positive Side (Source Area)', level=3)
        doc.add_paragraph(f"Description: {pos.get('description', 'N/A')}")
        doc.add_paragraph(f"Details: {pos.get('details', 'N/A')}")
        
        # Thermal findings
        thermal = obs.get('thermal_findings', {})
        if thermal and thermal.get('hotspot_temp'):
            doc.add_heading('Thermal Analysis', level=3)
            thermal_table = doc.add_table(rows=4, cols=2)
            thermal_table.style = 'Light Shading Accent 1'
            thermal_data = [
                ('Hotspot Temperature', thermal.get('hotspot_temp', 'N/A')),
                ('Coldspot Temperature', thermal.get('coldspot_temp', 'N/A')),
                ('Temperature Differential', thermal.get('temperature_differential', 'N/A')),
                ('Interpretation', thermal.get('interpretation', 'N/A')),
            ]
            for i, (label, value) in enumerate(thermal_data):
                thermal_table.cell(i, 0).text = label
                thermal_table.cell(i, 1).text = str(value)
        
        doc.add_paragraph()  # spacing
    
    # ---- SECTION 4: ROOT CAUSES ----
    doc.add_page_break()
    doc.add_heading('Section 4: Probable Root Causes', level=1)
    
    for i, cause in enumerate(ddr_data.get('probable_root_causes', []), 1):
        doc.add_heading(f"Cause {i}: {cause.get('cause', 'N/A')}", level=2)
        doc.add_paragraph(cause.get('explanation', 'N/A'))
        areas = cause.get('affected_areas', [])
        if areas:
            doc.add_paragraph(f"Affected Areas: {', '.join(areas)}")
    
    # ---- SECTION 5: SEVERITY ----
    doc.add_heading('Section 5: Severity Assessment', level=1)
    
    severity_data = ddr_data.get('severity_assessment', [])
    if severity_data:
        sev_table = doc.add_table(rows=len(severity_data) + 1, cols=4)
        sev_table.style = 'Medium Shading 1 Accent 1'
        headers = ['Area', 'Severity', 'Reasoning', 'Urgency']
        for j, header in enumerate(headers):
            sev_table.cell(0, j).text = header
        for i, sev in enumerate(severity_data, 1):
            sev_table.cell(i, 0).text = str(sev.get('area', 'N/A'))
            sev_table.cell(i, 1).text = str(sev.get('severity', 'N/A'))
            sev_table.cell(i, 2).text = str(sev.get('reasoning', 'N/A'))
            sev_table.cell(i, 3).text = str(sev.get('urgency', 'N/A'))
    
    # ---- SECTION 6: RECOMMENDED ACTIONS ----
    doc.add_page_break()
    doc.add_heading('Section 6: Recommended Actions', level=1)
    
    for action in ddr_data.get('recommended_actions', []):
        doc.add_heading(
            f"Priority {action.get('priority', 'N/A')}: {action.get('action', 'N/A')}", 
            level=2
        )
        doc.add_paragraph(action.get('description', 'N/A'))
        areas = action.get('applicable_areas', [])
        if areas:
            doc.add_paragraph(f"Applicable Areas: {', '.join(areas)}")
        doc.add_paragraph(f"Urgency: {action.get('estimated_urgency', 'N/A')}")
    
    # ---- SECTION 7: THERMAL ANALYSIS ----
    doc.add_page_break()
    doc.add_heading('Section 7: Thermal Analysis Summary', level=1)
    
    thermal_summary = ddr_data.get('thermal_analysis_summary', [])
    if thermal_summary:
        for ts in thermal_summary[:10]:  # Limit to first 10
            doc.add_heading(
                f"Thermal Image - Page {ts.get('thermal_page', 'N/A')}", level=2
            )
            t_table = doc.add_table(rows=5, cols=2)
            t_table.style = 'Light Shading Accent 1'
            t_data = [
                ('Hotspot', ts.get('hotspot', 'N/A')),
                ('Coldspot', ts.get('coldspot', 'N/A')),
                ('Differential', ts.get('differential', 'N/A')),
                ('Mapped Area', ts.get('area_mapped_to', 'N/A')),
                ('Finding', ts.get('finding', 'N/A')),
            ]
            for i, (label, value) in enumerate(t_data):
                t_table.cell(i, 0).text = label
                t_table.cell(i, 1).text = str(value)
            
            # Add thermal images
            therm_keys = list(therm_img_paths.keys())
            page_num = ts.get('thermal_page', 0)
            matching = [k for k in therm_keys if f"page{page_num}_" in k]
            for mk in matching[:2]:
                try:
                    doc.add_picture(therm_img_paths[mk], width=Inches(3))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    pass
            doc.add_paragraph()
    
    # ---- SECTION 8: SUMMARY TABLE ----
    doc.add_page_break()
    doc.add_heading('Section 8: Summary Table', level=1)
    
    summary_table_data = ddr_data.get('summary_table', [])
    if summary_table_data:
        s_table = doc.add_table(rows=len(summary_table_data) + 1, cols=4)
        s_table.style = 'Medium Shading 1 Accent 1'
        s_headers = ['Point No', 'Impacted Area (-ve side)', 'Point No', 'Exposed Area (+ve side)']
        for j, header in enumerate(s_headers):
            s_table.cell(0, j).text = header
        for i, row in enumerate(summary_table_data, 1):
            s_table.cell(i, 0).text = str(row.get('point_no', ''))
            s_table.cell(i, 1).text = str(row.get('impacted_area_negative', ''))
            s_table.cell(i, 2).text = str(row.get('point_no_positive', ''))
            s_table.cell(i, 3).text = str(row.get('exposed_area_positive', ''))
    
    # ---- SECTION 9: ADDITIONAL NOTES ----
    doc.add_page_break()
    doc.add_heading('Section 9: Additional Notes', level=1)
    for note in ddr_data.get('additional_notes', ['No additional notes.']):
        doc.add_paragraph(f"• {note}")
    
    # ---- SECTION 10: MISSING INFO ----
    doc.add_heading('Section 10: Missing or Unclear Information', level=1)
    for item in ddr_data.get('missing_or_unclear_information', ['None identified.']):
        doc.add_paragraph(f"• {item}")
    
    # ---- APPENDIX: IMAGES ----
    doc.add_page_break()
    doc.add_heading('Appendix A: Inspection Photographs', level=1)
    
    count = 0
    for img_id, img_path in insp_img_paths.items():
        if count >= 30:
            break
        try:
            doc.add_picture(img_path, width=Inches(3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Photo: {img_id}", style='Caption')
            count += 1
        except:
            continue
    
    doc.add_page_break()
    doc.add_heading('Appendix B: Thermal Images', level=1)
    
    count = 0
    for img_id, img_path in therm_img_paths.items():
        if count >= 30:
            break
        try:
            doc.add_picture(img_path, width=Inches(3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Thermal: {img_id}", style='Caption')
            count += 1
        except:
            continue
    
    # ---- LEGAL DISCLAIMER ----
    doc.add_page_break()
    doc.add_heading('Legal Disclaimer', level=1)
    legal = (
        "Information provided in this report is a general overview of the most obvious repairs that may be needed. "
        "It is not intended to be an exhaustive list. The inspection is not technically exhaustive; the property "
        "inspection provides the client with a basic overview of the condition of the unit. If such work is beyond "
        "the scope of the inspection, the inspector strongly recommends that the client consults a qualified "
        "Licensed Contractor Professional or Consulting Engineer."
    )
    doc.add_paragraph(legal)
    
    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────────
# REPORT GENERATION - PDF
# ─────────────────────────────────────────────

def generate_pdf(ddr_data, inspection_images, thermal_images, img_folder, output_path):
    """Generate a professional PDF report from DDR data."""
    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'DDRTitle', parent=styles['Title'],
        fontSize=24, textColor=HexColor('#8B0000'),
        spaceAfter=30, alignment=TA_CENTER
    )
    heading1 = ParagraphStyle(
        'DDRHeading1', parent=styles['Heading1'],
        fontSize=16, textColor=HexColor('#1a1a2e'),
        spaceBefore=20, spaceAfter=12
    )
    heading2 = ParagraphStyle(
        'DDRHeading2', parent=styles['Heading2'],
        fontSize=13, textColor=HexColor('#16213e'),
        spaceBefore=14, spaceAfter=8
    )
    heading3 = ParagraphStyle(
        'DDRHeading3', parent=styles['Heading3'],
        fontSize=11, textColor=HexColor('#0f3460'),
        spaceBefore=10, spaceAfter=6
    )
    body_style = ParagraphStyle(
        'DDRBody', parent=styles['Normal'],
        fontSize=10, leading=14, alignment=TA_JUSTIFY,
        spaceAfter=8
    )
    
    elements = []
    meta = ddr_data.get('report_metadata', {})
    
    # Cover
    elements.append(Spacer(1, 3*inch))
    elements.append(Paragraph('DETAILED DIAGNOSTIC REPORT', title_style))
    elements.append(Spacer(1, 0.3*inch))
    elements.append(Paragraph(meta.get('prepared_for', ''), ParagraphStyle('sub', parent=styles['Normal'], fontSize=14, alignment=TA_CENTER, textColor=HexColor('#555555'))))
    elements.append(Spacer(1, 0.2*inch))
    elements.append(Paragraph(f"Inspection Date: {meta.get('inspection_date', 'N/A')}", ParagraphStyle('info', parent=styles['Normal'], fontSize=11, alignment=TA_CENTER)))
    elements.append(Paragraph(f"Inspected By: {meta.get('inspected_by', 'N/A')}", ParagraphStyle('info2', parent=styles['Normal'], fontSize=11, alignment=TA_CENTER)))
    elements.append(PageBreak())
    
    # Property Info
    elements.append(Paragraph('Section 1: Property Information', heading1))
    info_data = [
        ['Parameter', 'Details'],
        ['Property Type', meta.get('property_type', 'N/A')],
        ['Address', meta.get('property_address', 'N/A')],
        ['Floors', str(meta.get('floors', 'N/A'))],
        ['Property Age', meta.get('property_age', 'N/A')],
        ['Previous Audit', meta.get('previous_structural_audit', 'N/A')],
        ['Previous Repair', meta.get('previous_repair_work', 'N/A')],
    ]
    t = Table(info_data, colWidths=[3*inch, 4*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#1a1a2e')),
        ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#ffffff')),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#cccccc')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#f8f9fa'), HexColor('#ffffff')]),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 0.3*inch))
    
    # Issue Summary
    elements.append(Paragraph('Section 2: Property Issue Summary', heading1))
    summary = ddr_data.get('property_issue_summary', 'Not Available')
    for para in summary.split('\n'):
        if para.strip():
            elements.append(Paragraph(para.strip(), body_style))
    
    elements.append(PageBreak())
    
    # Area-wise Observations
    elements.append(Paragraph('Section 3: Area-wise Observations', heading1))
    
    for obs in ddr_data.get('area_wise_observations', []):
        elements.append(Paragraph(f"Area {obs.get('area_id', '')}: {obs.get('area_name', '')}", heading2))
        
        neg = obs.get('negative_side', {})
        elements.append(Paragraph('Negative Side (Affected Area)', heading3))
        elements.append(Paragraph(f"<b>Description:</b> {neg.get('description', 'N/A')}", body_style))
        elements.append(Paragraph(f"<b>Details:</b> {neg.get('details', 'N/A')}", body_style))
        
        pos = obs.get('positive_side', {})
        elements.append(Paragraph('Positive Side (Source Area)', heading3))
        elements.append(Paragraph(f"<b>Description:</b> {pos.get('description', 'N/A')}", body_style))
        elements.append(Paragraph(f"<b>Details:</b> {pos.get('details', 'N/A')}", body_style))
        
        thermal = obs.get('thermal_findings', {})
        if thermal and thermal.get('hotspot_temp'):
            elements.append(Paragraph('Thermal Findings', heading3))
            th_data = [
                ['Hotspot', str(thermal.get('hotspot_temp', 'N/A'))],
                ['Coldspot', str(thermal.get('coldspot_temp', 'N/A'))],
                ['Differential', str(thermal.get('temperature_differential', 'N/A'))],
                ['Interpretation', str(thermal.get('interpretation', 'N/A'))],
            ]
            th_table = Table(th_data, colWidths=[2*inch, 5*inch])
            th_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#cccccc')),
                ('BACKGROUND', (0, 0), (0, -1), HexColor('#e8f4fd')),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            elements.append(th_table)
        
        elements.append(Spacer(1, 0.2*inch))
    
    # Root Causes
    elements.append(PageBreak())
    elements.append(Paragraph('Section 4: Probable Root Causes', heading1))
    for i, cause in enumerate(ddr_data.get('probable_root_causes', []), 1):
        elements.append(Paragraph(f"Cause {i}: {cause.get('cause', 'N/A')}", heading2))
        elements.append(Paragraph(cause.get('explanation', 'N/A'), body_style))
        areas = cause.get('affected_areas', [])
        if areas:
            elements.append(Paragraph(f"<b>Affected Areas:</b> {', '.join(areas)}", body_style))
    
    # Severity
    elements.append(Paragraph('Section 5: Severity Assessment', heading1))
    sev_data = ddr_data.get('severity_assessment', [])
    if sev_data:
        sev_rows = [['Area', 'Severity', 'Reasoning', 'Urgency']]
        for s in sev_data:
            sev_rows.append([
                str(s.get('area', '')),
                str(s.get('severity', '')),
                str(s.get('reasoning', ''))[:80],
                str(s.get('urgency', ''))
            ])
        sev_t = Table(sev_rows, colWidths=[1.5*inch, 1*inch, 3*inch, 1.2*inch])
        sev_t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#1a1a2e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#ffffff')),
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#cccccc')),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#f8f9fa'), HexColor('#ffffff')]),
        ]))
        elements.append(sev_t)
    
    # Recommended Actions
    elements.append(PageBreak())
    elements.append(Paragraph('Section 6: Recommended Actions', heading1))
    for action in ddr_data.get('recommended_actions', []):
        elements.append(Paragraph(
            f"Priority {action.get('priority', '')}: {action.get('action', 'N/A')}", heading2
        ))
        elements.append(Paragraph(action.get('description', 'N/A'), body_style))
        elements.append(Paragraph(f"<b>Urgency:</b> {action.get('estimated_urgency', 'N/A')}", body_style))
    
    # Summary Table
    elements.append(PageBreak())
    elements.append(Paragraph('Section 8: Summary Table', heading1))
    st_data = ddr_data.get('summary_table', [])
    if st_data:
        st_rows = [['#', 'Impacted Area (-ve)', '#', 'Exposed Area (+ve)']]
        for row in st_data:
            st_rows.append([
                str(row.get('point_no', '')),
                str(row.get('impacted_area_negative', '')),
                str(row.get('point_no_positive', '')),
                str(row.get('exposed_area_positive', ''))
            ])
        st_t = Table(st_rows, colWidths=[0.5*inch, 3*inch, 0.5*inch, 3*inch])
        st_t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#1a1a2e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#ffffff')),
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#cccccc')),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#f8f9fa'), HexColor('#ffffff')]),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        elements.append(st_t)
    
    # Additional Notes
    elements.append(Spacer(1, 0.3*inch))
    elements.append(Paragraph('Section 9: Additional Notes', heading1))
    for note in ddr_data.get('additional_notes', ['No additional notes.']):
        elements.append(Paragraph(f"• {note}", body_style))
    
    # Missing Info
    elements.append(Paragraph('Section 10: Missing or Unclear Information', heading1))
    for item in ddr_data.get('missing_or_unclear_information', ['None identified.']):
        elements.append(Paragraph(f"• {item}", body_style))
    
    doc.build(elements)
    return output_path


# ─────────────────────────────────────────────
# FLASK ROUTES
# ─────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/generate', methods=['POST'])
def generate():
    """Main endpoint to process uploads and generate DDR."""
    try:
        # Validate inputs
        api_key = request.form.get('api_key', '').strip()
        if not api_key:
            return jsonify({'error': 'Google Gemini API key is required'}), 400
        
        inspection_file = request.files.get('inspection_report')
        thermal_file = request.files.get('thermal_report')
        
        if not inspection_file or not thermal_file:
            return jsonify({'error': 'Both Inspection Report and Thermal Report PDFs are required'}), 400
        
        if not (allowed_file(inspection_file.filename) and allowed_file(thermal_file.filename)):
            return jsonify({'error': 'Only PDF files are accepted'}), 400
        
        # Save uploaded files
        session_id = str(uuid.uuid4())[:8]
        session_folder = os.path.join(app.config['UPLOAD_FOLDER'], session_id)
        os.makedirs(session_folder, exist_ok=True)
        
        insp_path = os.path.join(session_folder, 'inspection.pdf')
        therm_path = os.path.join(session_folder, 'thermal.pdf')
        inspection_file.save(insp_path)
        thermal_file.save(therm_path)
        
        # Extract data from PDFs
        insp_text, insp_images, insp_pages = extract_from_pdf(insp_path)
        therm_text, therm_images, therm_pages = extract_from_pdf(therm_path)
        
        # Generate DDR using Gemini
        ddr_data = generate_ddr_with_gemini(
            api_key, insp_text, therm_text,
            len(insp_images), len(therm_images)
        )
        
        # Create output files
        output_folder = os.path.join(app.config['OUTPUT_FOLDER'], session_id)
        os.makedirs(output_folder, exist_ok=True)
        img_folder = os.path.join(output_folder, 'images')
        os.makedirs(img_folder, exist_ok=True)
        
        docx_path = os.path.join(output_folder, 'DDR_Report.docx')
        pdf_path = os.path.join(output_folder, 'DDR_Report.pdf')
        json_path = os.path.join(output_folder, 'DDR_Data.json')
        
        # Save JSON
        with open(json_path, 'w') as f:
            json.dump(ddr_data, f, indent=2)
        
        # Generate DOCX
        generate_docx(ddr_data, insp_images, therm_images, img_folder, docx_path)
        
        # Generate PDF
        generate_pdf(ddr_data, insp_images, therm_images, img_folder, pdf_path)
        
        # Prepare preview data
        preview = {
            'session_id': session_id,
            'metadata': ddr_data.get('report_metadata', {}),
            'summary': ddr_data.get('property_issue_summary', ''),
            'observations': ddr_data.get('area_wise_observations', []),
            'root_causes': ddr_data.get('probable_root_causes', []),
            'severity': ddr_data.get('severity_assessment', []),
            'actions': ddr_data.get('recommended_actions', []),
            'additional_notes': ddr_data.get('additional_notes', []),
            'missing_info': ddr_data.get('missing_or_unclear_information', []),
            'summary_table': ddr_data.get('summary_table', []),
            'thermal_summary': ddr_data.get('thermal_analysis_summary', []),
            'num_inspection_images': len(insp_images),
            'num_thermal_images': len(therm_images),
            'inspection_images_b64': [{'id': img['id'], 'b64': img['b64'], 'ext': img['ext']} for img in insp_images[:20]],
            'thermal_images_b64': [{'id': img['id'], 'b64': img['b64'], 'ext': img['ext']} for img in therm_images[:20]],
        }
        
        return jsonify({'success': True, 'data': preview})
    
    except json.JSONDecodeError as e:
        return jsonify({'error': f'AI response parsing error: {str(e)}. Try again.'}), 500
    except Exception as e:
        return jsonify({'error': f'Processing error: {str(e)}'}), 500


@app.route('/download/<session_id>/<file_type>')
def download(session_id, file_type):
    """Download generated report files."""
    output_folder = os.path.join(app.config['OUTPUT_FOLDER'], session_id)
    
    if file_type == 'docx':
        filepath = os.path.join(output_folder, 'DDR_Report.docx')
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif file_type == 'pdf':
        filepath = os.path.join(output_folder, 'DDR_Report.pdf')
        mimetype = 'application/pdf'
    elif file_type == 'json':
        filepath = os.path.join(output_folder, 'DDR_Data.json')
        mimetype = 'application/json'
    else:
        return jsonify({'error': 'Invalid file type'}), 400
    
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found. Please generate the report first.'}), 404
    
    return send_file(filepath, as_attachment=True, mimetype=mimetype)


if __name__ == '__main__':
    os.makedirs('templates', exist_ok=True)
    app.run(debug=True, host='0.0.0.0', port=5000)
